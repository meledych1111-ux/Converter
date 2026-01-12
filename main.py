# main.py — DocLens Lite: полная версия без ошибок
import os
import tempfile
from flask import Flask, request, send_file, render_template_string
from pdf2image import convert_from_path
from pdf2image.pdfinfo import pdfinfo_from_path
import pytesseract
from PIL import Image
import cv2
import numpy as np
import camelot
import pandas as pd
from docx import Document
from weasyprint import HTML
from datetime import datetime
from sqlalchemy import create_engine, Column, Integer, String, DateTime
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker

app = Flask(__name__)
Base = declarative_base()
engine = create_engine("sqlite:///history.db")
SessionLocal = sessionmaker(bind=engine)

class Log(Base):
    __tablename__ = 'logs'
    id = Column(Integer, primary_key=True)
    filename = Column(String)
    result = Column(String)  # 'text', 'table'
    format = Column(String)  # 'docx', 'xlsx', 'pdf'
    ts = Column(DateTime, default=datetime.utcnow)

Base.metadata.create_all(engine)

def ocr_image(img):
    """Распознать текст с изображения."""
    img_cv = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)
    gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
    _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    return pytesseract.image_to_string(thresh, config=r'--oem 3 --psm 6 -l eng+rus+chi_sim')

def extract_tables_from_pdf(pdf_path, page_num):
    """Попытаться извлечь таблицы со страницы PDF."""
    tables = camelot.read_pdf(pdf_path, pages=str(page_num), flavor='lattice')
    if not tables:
        tables = camelot.read_pdf(pdf_path, pages=str(page_num), flavor='stream')
    return [t.df for t in tables if t.shape[0] > 1]

def get_pages(file_path, is_pdf, max_pages=10):
    """Получить список страниц (ограничено для скорости)."""
    if not is_pdf:
        return [Image.open(file_path).convert('RGB')]
    info = pdfinfo_from_path(file_path)
    total = min(int(info["Pages"]), max_pages)
    return convert_from_path(file_path, first_page=1, last_page=total, dpi=150)

@app.route('/')
def index():
    return render_template_string('''
<!doctype html>
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>DocLens</title>
<style>
body {
    font-family: -apple-system, BlinkMacSystemFont, sans-serif;
    padding: 20px;
    background: #f9f9f9;
    margin: 0;
}
h2 {
    color: #1e88e5;
    text-align: center;
}
input, button {
    width: 100%;
    padding: 14px;
    margin: 8px 0;
    border: 1px solid #ccc;
    border-radius: 10px;
    box-sizing: border-box;
}
button {
    background: #1e88e5;
    color: white;
    font-weight: bold;
    border: none;
}
.buttons {
    display: grid;
    gap: 8px;
}
</style>
<h2>📄 DocLens Lite</h2>
<form method=post enctype=multipart/form-data>
  <input type=file name=file accept=".pdf,.jpg,.jpeg,.png" required>
  <div class="buttons">
    <button type=submit name=format value=docx>📥 Word (.docx)</button>
    <button type=submit name=format value=xlsx>📊 Excel (.xlsx)</button>
    <button type=submit name=format value=pdf>🖨 PDF</button>
  </div>
</form>
''')

@app.route('/process', methods=['POST'])
def process():
    file = request.files['file']
    fmt = request.form['format']
    original_filename = file.filename
    is_pdf = original_filename.lower().endswith('.pdf')
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(original_filename)[1]) as tmp:
        file.save(tmp.name)
        input_path = tmp.name

    session = SessionLocal()
    output_path = None
    has_tables = False

    try:
        pages = get_pages(input_path, is_pdf, max_pages=10)
        results = []

        # Обработка каждой страницы
        for i, img in enumerate(pages, 1):
            if is_pdf:
                tables = extract_tables_from_pdf(input_path, i)
                if tables:
                    has_tables = True
                    results.append(('table', tables))
                else:
                    text = ocr_image(img)
                    results.append(('text', text))
            else:
                text = ocr_image(img)
                results.append(('text', text))

        # Экспорт
        if fmt == 'xlsx':
            if not has_tables:
                return '<h3 style="color:red; text-align:center;">❌ Таблицы не найдены</h3>'
            writer = pd.ExcelWriter('/tmp/out.xlsx', engine='openpyxl')
            sheet_index = 1
            for typ, data in results:
                if typ == 'table':
                    for df in data:
                        sheet_name = f'Sheet{sheet_index}'
                        df.to_excel(writer, sheet_name=sheet_name, index=False)
                        sheet_index += 1
            writer.close()
            output_path = '/tmp/out.xlsx'
            download_name = 'tables.xlsx'

        elif fmt == 'docx':
            doc = Document()
            doc.add_heading('Результат обработки', 0)
            for i, (typ, data) in enumerate(results, 1):
                doc.add_heading(f'Страница {i}', level=2)
                if typ == 'table':
                    for df in data:
                        t = doc.add_table(rows=df.shape[0], cols=df.shape[1])
                        t.style = 'Table Grid'
                        for r in range(df.shape[0]):
                            for c in range(df.shape[1]):
                                cell_text = str(df.iloc[r, c])[:500]
                                t.cell(r, c).text = cell_text
                else:
                    doc.add_paragraph(data if data.strip() else "[Текст не распознан]")
            output_path = '/tmp/out.docx'
            doc.save(output_path)
            download_name = 'result.docx'

        else:  # PDF
            html = '''
            <!DOCTYPE html>
            <html><head><meta charset="utf-8">
            <style>
                body { font-family: Arial, sans-serif; margin: 40px; }
                h1 { color: #333; }
                h2 { margin-top: 30px; border-bottom: 1px solid #ccc; padding-bottom: 5px; }
                table { border-collapse: collapse; width: 100%; margin: 10px 0; }
                th, td { border: 1px solid #999; padding: 6px; text-align: left; }
                pre { white-space: pre-wrap; background: #f9f9f9; padding: 10px; border-radius: 4px; }
            </style>
            </head><body>
            <h1>Результат обработки</h1>
            '''
            for i, (typ, data) in enumerate(results, 1):
                html += f'<h2>Страница {i}</h2>'
                if typ == 'table':
                    for df in data:
                        html += '<table>'
                        for _, row in df.iterrows():
                            html += '<tr>' + ''.join(f'<td>{cell}</td>' for cell in row) + '</tr>'
                        html += '</table>'
                else:
                    escaped = (data or "[Текст не распознан]").replace("&", "&amp;").replace("<", "&lt;")
                    html += f'<pre>{escaped}</pre>'
            html += '</body></html>'
            output_path = '/tmp/out.pdf'
            HTML(string=html).write_pdf(output_path)
            download_name = 'result.pdf'

        # Логируем
        session.add(Log(filename=original_filename, result='table' if has_tables else 'text', format=fmt))
        session.commit()

        return send_file(output_path, as_attachment=True, download_name=download_name)

    except Exception as e:
        return f'<h3 style="color:red; text-align:center;">Ошибка: {str(e)[:200]}</h3>'
    finally:
        session.close()
        for p in [input_path, output_path]:
            if p and os.path.exists(p):
                try:
                    os.remove(p)
                except:
                    pass

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=3000)
